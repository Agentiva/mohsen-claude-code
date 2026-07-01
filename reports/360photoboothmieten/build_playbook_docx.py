# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INK = RGBColor(0x0A, 0x0A, 0x0F)
GRN = RGBColor(0x1F, 0x8A, 0x5B)
MUT = RGBColor(0x6B, 0x6B, 0x75)
BLU = RGBColor(0x2A, 0x6F, 0xFD)

doc = Document()
st = doc.styles['Normal']
st.font.name = 'Calibri'; st.font.size = Pt(11)
st.paragraph_format.space_after = Pt(6); st.paragraph_format.line_spacing = 1.15

def shade(par, hexfill):
    pPr = par._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexfill)
    pPr.append(shd)

def title(text):
    p = doc.add_paragraph(); r = p.add_run(text); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = INK
    return p

def subtitle(text):
    p = doc.add_paragraph(); r = p.add_run(text); r.font.size = Pt(12); r.font.color.rgb = MUT
    return p

def block_header(text, fill='0A0A0F'):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(14)
    r = p.add_run('  ' + text + '  '); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    shade(p, fill); return p

def para(text, italic=False, color=None, size=11, bold=False):
    p = doc.add_paragraph(); r = p.add_run(text); r.italic = italic; r.bold = bold; r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return p

def kv(label, value):
    p = doc.add_paragraph(); r = p.add_run(label); r.bold = True; p.add_run(value); return p

def persona(name, titles, pains):
    p = doc.add_paragraph(); r = p.add_run(name); r.bold = True; r.font.size = Pt(13)
    p.paragraph_format.space_before = Pt(8)
    para(titles, italic=True, color=MUT)
    hp = doc.add_paragraph(); rr = hp.add_run('Pain Points:'); rr.bold = True
    for b in pains:
        bp = doc.add_paragraph(b, style='List Bullet'); bp.paragraph_format.space_after = Pt(2)

def named(titletext, body):
    p = doc.add_paragraph(); r = p.add_run(titletext); r.bold = True; r.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    para(body)

# ============ HEADER ============
title('Playbook: Modulare B2B Brand Activations für skalierende Unternehmen')
subtitle('360° Company / 360°PhotoBoothMieten (Threesixtymedia) · Kampagnen-Playbook · Sprache: de · erstellt von amplifa')
n = doc.add_paragraph()
nr = n.add_run('Repositionierung (GTM): Angepasst auf die freigegebenen First-Touch-Mails (Beispiel Parloa/Series D). '
               'Angebot wird als modularer B2B-Marketing-Partner geführt (Brand Activations, LED-Messesysteme, interaktive '
               'Eventmodule, Medienproduktion) — die 360° Video Booth ist das Flaggschiff-Modul. Zielgruppe: skalierende, '
               'oft finanzierte B2B-/Tech-Unternehmen mit Messe-, Partner-Event- und Roadshow-Bedarf. Signal-getriebene '
               'Ansprache. Kein Fathom-Onboarding vorhanden; Unbelegtes mit „(zu verifizieren)" markiert.')
nr.italic = True; nr.font.size = Pt(9.5); nr.font.color.rgb = MUT

# ============ BLOCK 1 ============
block_header('Block 1 — Product Description')
para('Die 360° Company (rechtlich Threesixtymedia, Marke „360°PhotoBoothMieten", Sitz Stolberg/NRW, Inhaber Youssef Ismail) '
     'ist ein B2B-Marketing-Partner für modulare Brand Activations. Das Unternehmen unterstützt skalierende Unternehmen dabei, '
     'Marke, Produkt und Zielgruppe an physischen Touchpoints erlebbar zu machen — auf Messen, Partner-Events, Roadshows und '
     'Enterprise-Events. Das Portfolio umfasst modulare Brand Activations, LED-Messesysteme, interaktive Eventmodule sowie '
     'hochwertige Medienproduktion. Die genaue Unternehmensgröße wird auf der Website nicht genannt; gearbeitet wird '
     'projektbasiert mit Personal vor Ort, der Auftritt ist zweisprachig (DE/EN) mit Schwerpunkt West-DACH.')
para('Ein zentrales Modul ist die 360° Video Booth: Sie verbindet Marke und Zielgruppe an einem physischen Touchpoint mit '
     'direkt nutzbarem Social Content, hoher Interaktion und optionaler Lead-Erfassung. Technisch fasst die Plattform bis zu '
     '5 Personen gleichzeitig, filmt mit 120 fps und erzeugt gebrandete Clips von bis zu 25 Sekunden, die per QR-Code sofort '
     'zum Teilen bereitstehen; Live-Gadgets, Event-Assistent sowie Auf- und Abbau sind Teil des schlüsselfertigen Setups. '
     'Der Booth wird vollständig im Corporate Design des Kunden gebrandet und lässt sich als eines von mehreren '
     'Aktivierungsmodulen in einen größeren Messe- oder Eventauftritt einbetten.')
para('Im Markt positioniert sich die 360° Company als Partner für messbares Live-Marketing statt austauschbarem „Fotospaß": '
     'ein Ansprechpartner, der Markenpräsenz, Social Content, Reichweite und qualifizierte Interaktion an realen Touchpoints '
     'zusammenführt. Besonders relevant wird das in Skalierungsphasen — nach Finanzierungsrunden, beim Ausbau des '
     'Partner-Ökosystems und bei internationaler Expansion — wenn die Qualität physischer Touchpoints (Messen, Partnerformate, '
     'Roadshows) mit dem Wachstum Schritt halten muss.')
kv('INDUSTRY: ', 'B2B Brand Activation, Live-Marketing, Experiential Marketing, Messe-/Eventservice, LED-Messesysteme, '
                 'Medienproduktion, Eventmodule')
kv('USPs: ', '1. Modularer B2B-Marketing-Partner: Brand Activations, LED-Messesysteme, interaktive Eventmodule und '
             'Medienproduktion aus einer Hand. 2. 360° Video Booth als Flaggschiff-Modul für Social Content, Reichweite und '
             'Interaktion am physischen Touchpoint. 3. Schlüsselfertig inklusive Personal, Auf-/Abbau und Anfahrt — ein '
             'Ansprechpartner, kein Organisationsaufwand. 4. Vollständiges Branding im Corporate Design für einen konsistenten '
             'Markenauftritt auf jedem Touchpoint. 5. Bis zu 5 Personen, 120 fps und Clips bis 25 Sekunden mit sofortigem '
             'QR-Sharing. 6. Optionale, (zu verifizieren) DSGVO-konforme Lead-Erfassung direkt an der Aktivierung. '
             '7. Ausrichtung auf skalierende B2B-Unternehmen mit Messe-, Partner-Event- und Roadshow-Bedarf. '
             '8. Kombinierbar zu einem Gesamtauftritt (LED + Booth + Medienproduktion) statt Einzel-Gimmick.')

# ============ BLOCK 2 ============
block_header('Block 2 — Value Proposition')
para('Die 360° Company hilft skalierenden B2B-Unternehmen, ihre Marke an physischen Touchpoints — Messen, Partner-Events und '
     'Roadshows — auf Enterprise-Niveau erlebbar zu machen und daraus direkt Social Content, Reichweite und qualifizierte '
     'Interaktion zu erzeugen. Über modulare Brand Activations, LED-Messesysteme, interaktive Eventmodule und Medienproduktion '
     'aus einer Hand — mit der 360° Video Booth als kompaktem Aktivierungsmodul und optionaler Lead-Erfassung — entsteht '
     'messbare Markenwirkung, ohne dass das eigene Team Aufwand oder Umsetzungsrisiko trägt.')

# ============ BLOCK 3 ============
block_header('Block 3 — Target Personas (5)')
persona('Head of Marketing Sarah (Scale-up)',
        'Head of Marketing, VP Marketing, CMO, Marketingleitung (skalierendes B2B-/Tech-Unternehmen)',
        ['Sie muss die Markenpräsenz mit dem Unternehmenswachstum skalieren — nach Funding und bei internationaler Expansion steigen die Erwartungen an jeden Auftritt.',
         'Sie will, dass physische Touchpoints (Messen, Roadshows) das gleiche Premium-Niveau haben wie die digitale Marke.',
         'Sie steht unter Druck, Marketing-Budget mit messbarer Wirkung — Reichweite, Interaktion, Leads — zu rechtfertigen.',
         'Sie braucht Partner, die schlüsselfertig liefern, damit ihr schlankes Team sich nicht in Eventlogistik verliert.',
         'Sie sucht Aktivierungen, die direkt teilbaren Social Content erzeugen und die Marke über das Event hinaus tragen.'])
persona('Head of Partnerships David',
        'Head of Partnerships, Partner Marketing Manager, Head of Alliances, Ecosystem Lead',
        ['Er baut ein wachsendes Partner-Ökosystem auf und muss Partner-Events und Co-Marketing-Formate hochwertig inszenieren.',
         'Er will, dass gemeinsame Auftritte mit Partnern Aufmerksamkeit erzeugen und die Marke bei deren Zielgruppen sichtbar machen.',
         'Er braucht Module, die sich flexibel in Partnerformate, Messen und Roadshows einbetten lassen.',
         'Er muss aus Partner-Events verwertbare Kontakte und Content mitnehmen, nicht nur Händeschütteln.',
         'Er hat selten eigene Eventressourcen und braucht einen verlässlichen Umsetzungspartner.'])
persona('Field- & Event-Marketing-Manager Nina',
        'Field Marketing Manager, Event Marketing Manager, Head of Events, Trade-Fair Manager',
        ['Sie verantwortet mehrere Messen, Roadshow-Stopps und Enterprise-Events pro Jahr und muss jeden Auftritt zum Highlight machen.',
         'Sie steht unter Druck, Standfrequenz, Verweildauer und qualifizierte Gespräche messbar zu steigern.',
         'Sie will ein Modul, das reibungslos funktioniert und inklusive Personal, Auf- und Abbau kommt.',
         'Sie braucht reproduzierbare Setups, die über mehrere Termine hinweg identisch hochwertig sind.',
         'Sie sucht ein frisches Format jenseits der Standard-Fotobox, das social-tauglichen Content erzeugt.'])
persona('Head of Brand Alex',
        'Head of Brand, Brand Manager, Brand Activation Manager, Head of Brand Experience',
        ['Er muss sicherstellen, dass die Marke an jedem physischen Touchpoint konsistent und premium auftritt.',
         'Er will Aktivierungen, die die Markenbotschaft erlebbar machen, nicht nur dekorieren.',
         'Er braucht hochwertigen, gebrandeten Content aus jedem Event für Always-on-Kanäle.',
         'Er sucht ein modulares Aktivierungskonzept, das mit der Marke mitwächst.'])
persona('Demand-Gen-Lead Tom',
        'Demand Generation Manager, Growth Marketing Lead, Head of Growth, Pipeline Marketing Manager',
        ['Er muss aus teuren Event-Investitionen belegbare Pipeline und Leads ziehen.',
         'Er will physische Touchpoints in messbare, DSGVO-konforme Kontakte übersetzen.',
         'Er braucht Formate, die hohe Interaktion erzeugen und gleichzeitig Daten liefern.',
         'Er steht unter Druck, den ROI jeder Aktivierung an Reichweite und Interaktion festzumachen.'])

# ============ BLOCK 4 ============
block_header('Block 4 — Use Cases (6)')
named('Markenpräsenz in der Skalierungsphase auf Enterprise-Niveau heben',
      'Nach einer Finanzierungsrunde oder bei internationaler Expansion steigen die Erwartungen an jeden physischen Auftritt, '
      'während das Marketingteam oft schlank bleibt. Die 360° Company liefert modulare Brand Activations inklusive '
      'LED-Messesystemen, Eventmodulen und Medienproduktion aus einer Hand. So halten Messen, Partnerformate und Roadshows '
      'qualitativ mit dem Wachstum Schritt, ohne dass intern Kapazität aufgebaut werden muss.')
named('Partner-Event- und Co-Marketing-Aktivierung',
      'Beim Ausbau des Partner-Ökosystems müssen gemeinsame Auftritte Aufmerksamkeit erzeugen und die Marke bei den Zielgruppen '
      'der Partner sichtbar machen. Die 360° Video Booth und weitere Module lassen sich flexibel in Partner-Events einbetten und '
      'erzeugen teilbaren, gebrandeten Content sowie qualifizierte Interaktion. Das Ergebnis ist ein Partnerformat, das Reichweite '
      'und Kontakte über den Event-Tag hinaus liefert.')
named('Enterprise-tauglicher Messeauftritt aus einem Guss',
      'Aussteller wollen einen Stand, der aus der Masse heraussticht und zum Premium-Markenbild passt. Die 360° Company kombiniert '
      'LED-Messesysteme, interaktive Eventmodule und die 360° Video Booth zu einem stimmigen Gesamtauftritt inklusive Personal. '
      'Das steigert Standfrequenz und Verweildauer und liefert gleichzeitig gebrandeten Social Content vom Messeauftritt.')
named('Reproduzierbare Roadshow über mehrere Standorte',
      'Bei internationalen Roadshows muss dieselbe Markeninszenierung an jedem Stopp verlässlich funktionieren. Die 360° Company '
      'stellt ein standardisiertes, gebrandetes Modul-Setup bereit, das an jedem Termin identisch hochwertige Aktivierung und '
      'Content produziert. Das Ergebnis ist ein konsistenter Markenauftritt und ein durchgängiger Content-Strom über die gesamte Tour.')
named('Produkt- oder Milestone-Launch in Social Content übersetzen',
      'Ein Produktlaunch oder ein Unternehmens-Meilenstein (z. B. Funding, neue Partnerschaft) braucht sichtbare, teilbare '
      'Momente. Die 360° Company inszeniert den Anlass mit Aktivierungsmodul und hochwertiger Medienproduktion, sodass daraus '
      'direkt nutzbarer Social Content und Reichweite entstehen. Das Ergebnis ist ein Launch, der über die Anwesenden hinaus wirkt.')
named('Qualifizierte Lead-Erfassung am physischen Touchpoint',
      'Demand-Gen- und Vertriebsverantwortliche brauchen aus Event-Investitionen verwertbare Kontakte statt reiner Sichtkontakte. '
      'Die 360° Video Booth lässt sich (zu verifizieren) mit einer Opt-in-Datenerfassung kombinieren, sodass Besucher ihren '
      'gebrandeten Clip gegen ihre Kontaktdaten erhalten. So wird aus hoher Interaktion eine messbare, DSGVO-konforme Kontaktliste.')

# ============ BLOCK 5 ============
block_header('Block 5 — Reference Customers (zu verifizieren)')
para('(zu verifizieren – aus Onboarding ergänzen) Es ließen sich keine belegten, namentlich nennbaren Referenzkunden mit '
     'Kundenstimme identifizieren. Im Kickoff 2–3 nennbare Referenzprojekte mit Marke, Format (Messe/Partner-Event/Roadshow), '
     'Ansprechpartner und – idealerweise – Kennzahlen (erstellte Clips, Reichweite, generierte Leads) sowie Logo-/Case-Freigabe '
     'einsammeln. Besonders wertvoll wären Referenzen aus dem skalierenden B2B-/Tech-Segment. Bis dahin Block leer lassen, nicht '
     'erfinden.', italic=True, color=MUT)

# ============ BLOCK 6 ============
block_header('Block 6 — Proof Points (6)')
named('Modulares Portfolio aus einer Hand',
      'Die 360° Company führt Brand Activations, LED-Messesysteme, interaktive Eventmodule und Medienproduktion zusammen. Kunden '
      'erhalten einen Ansprechpartner für einen stimmigen Gesamtauftritt statt vieler Einzeldienstleister — ein zentrales '
      'Argument für skalierende Teams mit begrenzter interner Kapazität.')
named('360° Video Booth: bis zu 5 Personen, 120 fps, Clips bis 25 Sekunden',
      'Das Flaggschiff-Modul ist technisch klar spezifiziert: Plattform für bis zu fünf Personen gleichzeitig, Aufnahme mit '
      '120 fps und gebrandete Clips von bis zu 25 Sekunden — Grundlage für hochwertigen, teilbaren Content statt einfacher Schnappschüsse.')
named('Vollständiges Branding im Corporate Design',
      'Jedes Modul wird auf das Corporate Design des Kunden abgestimmt — von Logo über Farben bis zur Botschaft. So transportiert '
      'jeder Touchpoint konsistent die Marke, auf dem Event wie in Social Media.')
named('Schlüsselfertig inklusive Personal, Auf-/Abbau und Anfahrt',
      'Das Angebot umfasst Personal vor Ort sowie Auf- und Abbau; die Anfahrt wird deutschlandweit beworben. Das entlastet '
      'schlanke Marketing- und Eventteams und sichert einen reibungslosen Ablauf am Eventtag.')
named('Sofortiges Teilen und optionale Lead-Erfassung',
      'Aufnahmen stehen per QR-Code sofort zum Teilen bereit und erzeugen organische, markengetriebene Reichweite; optional '
      'lässt sich (zu verifizieren) eine DSGVO-konforme Lead-Erfassung ergänzen, die Interaktion in messbare Kontakte übersetzt.')
named('Marktrückenwind für physische B2B-Aktivierungen',
      'Der Bedarf ist strukturell belegt: Experiential Marketing macht rund 28–35 % der Marketingbudgets aus, 97 % der '
      'B2B-Marketer halten Vor-Ort-Events für maßgeblich, und die deutsche Messewirtschaft zählte 2025 rund 190.000 Aussteller '
      'auf 304 Messen (AUMA). Gerade in Skalierungsphasen gewinnen hochwertige physische Touchpoints an Bedeutung.')

# ============ APPENDIX: MESSAGING & TRIGGER ============
doc.add_page_break()
block_header('Anhang (nicht für App-Felder) — Messaging-Framework & Trigger', fill='2A6FFD')
para('Dieser Anhang übersetzt die Positionierung in die konkrete Erstansprache. Grundlage sind die vier freigegebenen '
     'First-Touch-Varianten (Beispiel Parloa). Er dient dem Kampagnen-Setup und gehört nicht in die App-Playbook-Felder.',
     italic=True, color=MUT)

named('Messaging-Prinzip: Signal → Angebot → Modul → Soft-CTA',
      'Jede Mail öffnet mit einem firmenspezifischen Trigger (z. B. Funding, Partner-Ökosystem, internationale Skalierung), '
      'positioniert die 360° Company als modularen B2B-Marketing-Partner, nennt die 360° Video Booth als konkretes Beispiel-Modul '
      'und schließt mit einem niederschwelligen CTA (kurzer digitaler Austausch / 30 Minuten). Ton: strategisch, seriös, premium, knapp.')

hp = doc.add_paragraph(); r = hp.add_run('Trigger / Buying-Signals für die Personalisierung des Openers:'); r.bold = True
for b in ['Finanzierungsrunde / Series X / frisches Kapital (Skalierungssignal)',
          'Ausbau des Partner-Ökosystems, neue Allianzen oder Partnerprogramme',
          'Internationale Expansion / neue Märkte / neue Standorte',
          'Angekündigte Messeteilnahme, Roadshow oder Partner-/Enterprise-Event (mit Datum = stärkster Trigger)',
          'Produktlaunch, Rebranding oder Unternehmens-Meilenstein',
          'Hiring in Field-/Event-/Partner-Marketing oder Brand']:
    doc.add_paragraph(b, style='List Bullet').paragraph_format.space_after = Pt(2)

hp = doc.add_paragraph(); r = hp.add_run('Zielpersonen (Empfänger der Erstansprache):'); r.bold = True
para('Head of Marketing / VP Marketing / CMO, Head of Partnerships / Partner Marketing, Field-/Event-Marketing-Manager, '
     'Head of Brand — bei skalierenden, oft finanzierten B2B-/Tech-Unternehmen. Reihenfolge: zuerst Partnerships & '
     'Field-/Event-Marketing (konkretester Anlass), dann Marketing-/Brand-Leads.')

hp = doc.add_paragraph(); r = hp.add_run('Freigegebene First-Touch-Varianten (Referenz für die AI):'); r.bold = True
r.font.color.rgb = INK

variants = [
 ('Version 1 — strategisch & seriös',
  'Sehr geehrter Herr Shaw,\n\ndie Series-D-Finanzierung über 350 Mio. US-Dollar und der Ausbau des Partner-Ökosystems zeigen, '
  'dass Parloa in eine starke Skalierungsphase geht.\n\nWir unterstützen B2B-Unternehmen als Marketing-Partner mit modularen '
  'Brand Activations, LED-Messesystemen, interaktiven Eventmodulen und hochwertiger Medienproduktion.\n\nUnsere 360° Video Booth '
  'ist dabei ein Beispiel: Sie verbindet Marke, Produkt und Zielgruppe an physischen Touchpoints mit direkt nutzbarem Social '
  'Content, hoher Interaktion und optionaler Lead-Erfassung.\n\nGerne zeige ich Ihnen in einem kurzen digitalen Austausch, wie '
  'ein solches Aktivierungsmodul für Parloa aussehen könnte.\n\nBeste Grüße\nYoussef Ismail'),
 ('Version 2 — etwas selbstbewusster',
  'Sehr geehrter Herr Shaw,\n\nParloa skaliert sichtbar international — mit starker Finanzierung, wachsendem Partner-Ökosystem '
  'und zunehmender Präsenz an Enterprise-Touchpoints.\n\nGenau dort setzen wir an: als B2B-Marketing-Partner für modulare Brand '
  'Activations, LED-Messesysteme, interaktive Eventformate und Content-Produktion.\n\nDie 360° Video Booth ist eines unserer '
  'Module, um Marke oder Produkt vor Ort erlebbar zu machen und daraus direkt Social Content, Reichweite und qualifizierte '
  'Interaktion zu erzeugen.\n\nFalls physische Aktivierungen für Partner-Events, Messen oder Roadshows relevant sind, würde ich '
  'Ihnen gerne einen kompakten Ansatz zeigen.\n\nBeste Grüße\nYoussef Ismail'),
 ('Version 3 — Premium & klar',
  'Sehr geehrter Herr Shaw,\n\nmit der aktuellen Skalierungsphase von Parloa wird die Qualität physischer Touchpoints noch '
  'relevanter — besonders bei Messen, Partnerformaten und Roadshows.\n\nWir entwickeln modulare B2B Brand Activations mit '
  'LED-Messesystemen, interaktiven Eventmodulen und hochwertiger Medienproduktion.\n\nUnsere 360° Video Booth ist ein Beispiel '
  'dafür: ein kompaktes Aktivierungsmodul, das Markenpräsenz, Social Content und Reichweite direkt vor Ort miteinander '
  'verbindet.\n\nGerne zeige ich Ihnen in 30 Minuten, wie ein solches Modul für Parloa eingesetzt werden könnte.\n\nBeste Grüße\n'
  'Youssef Ismail'),
 ('Version 4 — sehr kurz & direkt',
  'Sehr geehrter Herr Shaw,\n\nParloa befindet sich sichtbar in einer starken Skalierungsphase — international, partnergetrieben '
  'und mit hoher Relevanz für Messen, Roadshows und Enterprise-Events.\n\nWir unterstützen B2B-Unternehmen mit modularen Brand '
  'Activations, LED-Messesystemen, interaktiven Eventmodulen und Medienproduktion.\n\nDie 360° Video Booth ist dabei eines '
  'unserer Tools, um Marke, Produkt und Zielgruppe vor Ort mit Reichweite, Social Content und Interaktion zu verbinden.\n\nGerne '
  'würde ich Ihnen kurz zeigen, wie ein solches Format für Parloa aussehen könnte.\n\nBeste Grüße\nYoussef Ismail'),
]
for head, body in variants:
    p = doc.add_paragraph(); r = p.add_run(head); r.bold = True; r.font.size = Pt(11.5); r.font.color.rgb = BLU
    p.paragraph_format.space_before = Pt(8)
    bp = doc.add_paragraph(body); bp.paragraph_format.left_indent = Pt(14)
    for run in bp.runs: run.font.size = Pt(10.5)

out = '/home/user/mohsen-claude-code/reports/360photoboothmieten/360PhotoBoothMieten Playbook.docx'
doc.save(out)
print('saved', out)
